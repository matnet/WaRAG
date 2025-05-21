# document_processor.py

import fitz  # PyMuPDF
import os
import re
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
import logging
import tempfile
import docx  # python-docx for DOCX files

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def extract_text_from_pdf(pdf_path, chunk_size=1000, chunk_overlap=100):
    """
    Enhanced function to extract text from PDF files and split into manageable chunks
    
    Args:
        pdf_path (str): Path to the PDF file
        chunk_size (int): Size of text chunks
        chunk_overlap (int): Overlap between chunks
        
    Returns:
        list: List of Document objects with text chunks and metadata
    """
    if not os.path.exists(pdf_path):
        logger.error(f"PDF file not found: {pdf_path}")
        raise FileNotFoundError(f"PDF file not found: {pdf_path}")
    
    try:
        logger.info(f"Opening PDF: {pdf_path}")
        doc = fitz.open(pdf_path)
        
        # Extract document metadata
        metadata = {
            "title": doc.metadata.get("title", ""),
            "author": doc.metadata.get("author", ""),
            "subject": doc.metadata.get("subject", ""),
            "creator": doc.metadata.get("creator", ""),
            "producer": doc.metadata.get("producer", ""),
            "total_pages": len(doc),
            "file_name": os.path.basename(pdf_path),
            "file_type": "pdf"
        }
        logger.info(f"PDF metadata: {metadata}")
        
        # Extract text from each page with improved formatting
        all_text = []
        for page_num, page in enumerate(doc):
            # Get text blocks with their bounding boxes
            text_blocks = page.get_text("blocks")
            
            # Sort blocks by vertical position (top to bottom)
            text_blocks.sort(key=lambda b: b[1])  # Sort by y-coordinate
            
            # Extract and join the text
            page_text = "\n".join([block[4] for block in text_blocks if isinstance(block[4], str)])
            
            # Clean up excessive whitespace
            page_text = re.sub(r'\s+', ' ', page_text).strip()
            
            if page_text:  # Only add non-empty pages
                all_text.append({
                    "page": page_num + 1,
                    "text": page_text
                })
                logger.debug(f"Extracted {len(page_text)} characters from page {page_num + 1}")
        
        # Close the PDF
        doc.close()
        
        if not all_text:
            logger.warning(f"No text extracted from PDF: {pdf_path}")
            
            # Try OCR as a fallback for scanned documents
            # This is a placeholder - implement proper OCR handling if needed
            logger.info("No text found, PDF might be scanned or image-based")
            return []
        
        return create_document_chunks(all_text, metadata, chunk_size, chunk_overlap)
    
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        raise

def extract_text_from_docx(docx_path, chunk_size=1000, chunk_overlap=100):
    """
    Extract text from DOCX files and split into manageable chunks
    
    Args:
        docx_path (str): Path to the DOCX file
        chunk_size (int): Size of text chunks
        chunk_overlap (int): Overlap between chunks
        
    Returns:
        list: List of Document objects with text chunks and metadata
    """
    if not os.path.exists(docx_path):
        logger.error(f"DOCX file not found: {docx_path}")
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")
    
    try:
        logger.info(f"Opening DOCX: {docx_path}")
        doc = docx.Document(docx_path)
        
        # Extract document metadata (more limited than PDF)
        metadata = {
            "title": doc.core_properties.title or "",
            "author": doc.core_properties.author or "",
            "subject": doc.core_properties.subject or "",
            "file_name": os.path.basename(docx_path),
            "file_type": "docx"
        }
        logger.info(f"DOCX metadata: {metadata}")
        
        # Extract text from each paragraph
        paragraphs_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs_text.append(para.text.strip())
        
        # Extract text from tables too
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs_text.append(cell.text.strip())
        
        if not paragraphs_text:
            logger.warning(f"No text extracted from DOCX: {docx_path}")
            return []
        
        # Process paragraphs into pages (approx 3000 chars per page)
        all_text = []
        current_page_text = ""
        current_page = 1
        
        for para in paragraphs_text:
            current_page_text += para + "\n\n"
            
            # When we reach ~3000 chars, consider it a new "page"
            if len(current_page_text) > 3000:
                all_text.append({
                    "page": current_page,
                    "text": current_page_text.strip()
                })
                current_page_text = ""
                current_page += 1
        
        # Add the last page if it's not empty
        if current_page_text.strip():
            all_text.append({
                "page": current_page,
                "text": current_page_text.strip()
            })
        
        return create_document_chunks(all_text, metadata, chunk_size, chunk_overlap)
    
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        raise

def create_document_chunks(all_text, metadata, chunk_size=1000, chunk_overlap=100):
    """
    Common chunking logic for both PDF and DOCX documents
    
    Args:
        all_text: List of dicts with page numbers and text
        metadata: Document metadata
        chunk_size: Size of text chunks
        chunk_overlap: Overlap between chunks
        
    Returns:
        list: List of Document objects
    """
    # Create a text splitter with appropriate settings
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        separators=["\n\n", "\n", ". ", " ", ""]
    )
    
    # Process all pages into a single document first
    combined_text = "\n\n".join([f"Page {page_info['page']}: {page_info['text']}" 
                                for page_info in all_text])
    
    # Split text into chunks
    chunks = text_splitter.split_text(combined_text)
    
    # Create Document objects
    documents = []
    for i, chunk in enumerate(chunks):
        # Try to identify which page this chunk belongs to
        page_match = re.match(r"Page (\d+):", chunk)
        page_num = int(page_match.group(1)) if page_match else 0
        
        # Clean up the chunk text
        clean_chunk = re.sub(r"Page \d+: ", "", chunk)
        
        doc = Document(
            page_content=clean_chunk,
            metadata={
                "page": page_num,
                "chunk": i + 1,
                "source": metadata["file_name"],
                **metadata
            }
        )
        documents.append(doc)
    
    logger.info(f"Extracted {len(documents)} text chunks from {len(all_text)} pages")
    return documents

def process_document_from_bytes(doc_bytes, filename, chunk_size=1000, chunk_overlap=100):
    """
    Process document data from bytes (useful for processing docs from WhatsApp)
    
    Args:
        doc_bytes (bytes): The document file as bytes
        filename (str): Original filename
        chunk_size (int): Size of text chunks
        chunk_overlap (int): Overlap between chunks
        
    Returns:
        list: List of Document objects
    """
    # Determine document type from filename
    file_ext = os.path.splitext(filename.lower())[1]
    
    # Create a temporary file
    with tempfile.NamedTemporaryFile(suffix=file_ext, delete=False) as temp_file:
        temp_path = temp_file.name
        temp_file.write(doc_bytes)
    
    try:
        # Process based on file type
        if file_ext == '.pdf':
            documents = extract_text_from_pdf(temp_path, chunk_size, chunk_overlap)
        elif file_ext in ['.docx', '.doc']:
            documents = extract_text_from_docx(temp_path, chunk_size, chunk_overlap)
        else:
            raise ValueError(f"Unsupported document type: {file_ext}")
            
        # Update source in metadata to original filename
        for doc in documents:
            doc.metadata["source"] = filename
            doc.metadata["file_name"] = filename
            
        return documents
    finally:
        # Clean up the temp file
        if os.path.exists(temp_path):
            os.remove(temp_path)

# Backward compatibility functions
def process_pdf_from_bytes(pdf_bytes, filename="document.pdf", chunk_size=1000, chunk_overlap=100):
    """Backward compatibility function for PDF processing"""
    return process_document_from_bytes(pdf_bytes, filename, chunk_size, chunk_overlap)
